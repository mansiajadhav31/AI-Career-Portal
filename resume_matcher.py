import PyPDF2
import docx
import re
import os

# Common tech skills database
COMMON_SKILLS = [
    'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'express',
    'django', 'flask', 'spring', 'spring boot', 'sql', 'mongodb', 'postgresql', 'mysql',
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn', 'keras',
    'html', 'css', 'typescript', 'php', 'ruby', 'c++', 'c#', 'swift', 'kotlin',
    'flutter', 'react native', 'graphql', 'rest api', 'restful', 'redis', 'elasticsearch',
    'hadoop', 'spark', 'tableau', 'power bi', 'agile', 'scrum', 'jira', 'confluence',
    'selenium', 'cypress', 'jest', 'pytest', 'unit test', 'integration test',
    'ci/cd', 'devops', 'linux', 'unix', 'bash', 'powershell', 'terraform', 'ansible',
    'excel', 'word', 'powerpoint', 'outlook', 'communication', 'leadership', 'teamwork'
]

def extract_resume_text(file_path):
    """Extract text from PDF or DOCX file"""
    text = ""
    
    print(f"📖 Extracting text from: {file_path}")
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return ""
    
    try:
        if file_path.lower().endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"📄 Number of pages: {len(pdf_reader.pages)}")
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"   Page {page_num + 1}: {len(page_text)} chars")
                    else:
                        print(f"   Page {page_num + 1}: No text found (scanned PDF?)")
        
        elif file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            print(f"📄 DOCX paragraphs: {len(doc.paragraphs)}")
        
        else:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            print(f"📄 Text file length: {len(text)} chars")
    
    except Exception as e:
        print(f"❌ Extraction error: {e}")
        return ""
    
    print(f"✅ Total extracted text: {len(text)} chars")
    return text.lower()

def extract_skills(text):
    """Extract skills from text"""
    found_skills = []
    text_lower = text.lower()
    
    for skill in COMMON_SKILLS:
        # Check for whole word match or partial match
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    # Remove duplicates while preserving order
    unique_skills = []
    for skill in found_skills:
        if skill not in unique_skills:
            unique_skills.append(skill)
    
    return unique_skills

def calculate_match_score(resume_text, job_description):
    """Calculate match score between resume and job description"""
    
    print("🔍 Calculating match score...")
    print(f"📝 Resume text length: {len(resume_text)} chars")
    print(f"📝 Job description length: {len(job_description)} chars")
    
    # If no resume text
    if not resume_text:
        print("⚠️ No resume text found")
        return 0, []
    
    # Extract skills from both
    resume_skills = extract_skills(resume_text)
    job_skills = extract_skills(job_description) if job_description else []
    
    print(f"📋 Resume skills found ({len(resume_skills)}): {resume_skills[:10]}")
    print(f"📋 Job skills found ({len(job_skills)}): {job_skills[:10]}")
    
    # If no job description provided
    if not job_description or len(job_description.strip()) < 10:
        print("⚠️ No valid job description, using default score based on resume quality")
        # Base score on number of skills found
        base_score = min(80, 40 + (len(resume_skills) * 3))
        return base_score, resume_skills[:5]
    
    # If no job skills found
    if len(job_skills) == 0:
        print("⚠️ No skills found in job description")
        return 65, resume_skills[:3] if resume_skills else ['General fit']
    
    # Calculate skill match
    matched_skills = [skill for skill in job_skills if skill in resume_skills]
    skill_score = (len(matched_skills) / len(job_skills)) * 100
    print(f"🎯 Skill match: {len(matched_skills)}/{len(job_skills)} = {skill_score:.1f}%")
    
    # Simple text similarity (keyword matching)
    resume_words = set(re.findall(r'\b\w+\b', resume_text.lower()))
    job_words = set(re.findall(r'\b\w+\b', job_description.lower()))
    
    # Remove common stop words
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                  'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
                  'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing'}
    
    resume_words = resume_words - stop_words
    job_words = job_words - stop_words
    
    if len(job_words) > 0:
        common_words = resume_words & job_words
        text_similarity = (len(common_words) / len(job_words)) * 100
    else:
        text_similarity = 0
    
    print(f"📝 Text similarity: {text_similarity:.1f}%")
    
    # Combined score (80% skill match, 20% text similarity)
    final_score = int((skill_score * 0.8) + (text_similarity * 0.2))
    final_score = max(0, min(100, final_score))
    
    print(f"🏆 Final match score: {final_score}%")
    print(f"✅ Matched skills: {matched_skills}")
    
    return final_score, matched_skills